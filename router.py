from fastapi import APIRouter, HTTPException, status, UploadFile,File,Form
from pydantic import BaseModel

from openai import AzureOpenAI
from docx import Document
from dotenv import load_dotenv
import uuid

load_dotenv()

router_ = APIRouter()

class Questions(BaseModel):
    ask: str


import os
import re

current_dir = os.path.dirname(os.path.abspath(__file__))


client = AzureOpenAI(
    api_version=os.getenv("OPENAI_API_VERSION"),
    azure_endpoint=os.getenv("ENDPOINT_DANI_"),
    api_key=os.getenv("API_KEY_DANI_"),
)

def get_id():
    return str(uuid.uuid4())

async def read_md(file_path: str):
    doc_dir = os.path.normpath(os.path.join(current_dir,file_path))
    with open(file=doc_dir, mode="r", encoding="utf-8") as file:
        prompt = file.read()
        return prompt
    
async def procesar_prospecto(md_text):

    processed_text = re.sub(r'~~(.*?)~~', r'<span style="color: red;">~~\1~~</span>', md_text)

    processed_text = re.sub(r'\*\*(.*?)\*\*', r'<span style="color: green;">**\1**</span>', processed_text)
    return processed_text
    




async def chat_completation_doc(md_text: str | None):

    system_prompt = """"
                        Recibiras un texto en formato de Markdown.
                        Quitar del texto MD los elementos "**".
                        Quitar del texto MD los elementos tachados con "~~".
                        Quitar encabezados desde la segunda página en adelante.
                        Quitar los números de página.
                        La salida debe ser texto Markdown unicamente, sin introducción ni conclusión."""
    
    if md_text is None or md_text == "":
        content_user = ""
    else:
        content_user = md_text
    
    response = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"{system_prompt}",
            },
            {
                "role": "user",
                "content": f"{content_user}",
            }
        ],
        temperature=0.5,
        tools=[],
        top_p=1.0,
        model="gpt-4.1"
    )
    respuesta = response.choices[0].message.content
    return respuesta



async def chat_comparation(prompt: str):
    response = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"{prompt}",
            },
            {
                "role": "user",
                "content": "",
            }
        ],
        #max_tokens=4096,
        temperature=0.5,
        tools=[],
        top_p=1.0,
        model="gpt-4.1"
    )
    respuesta = response.choices[0].message.content
    return respuesta

 

async def create_docx(parragragh: str):
    doc = Document()  

    doc.add_heading('Reporte Limpio', level=1)  

    doc.add_paragraph(f'{parragragh}')

    id = get_id()

    # Ruta al archivo dentro de la carpeta 'documents'
    file_name = f'reporte_isp_{id}.docx'
    file_path = os.path.join('documents', file_name)

    doc.save(file_path)
    print("Listo creado el documento docx")


    
@router_.get("/")
async def root():
    return {"message":"Bienvenido al Backend"}
    


from fastapi import UploadFile, File, APIRouter
from langchain.prompts import PromptTemplate
from views import TempleFileHannder
import asyncio  

@router_.post("/mini_rag")
async def comparar_docus(file_1: UploadFile = File(...), file_2: UploadFile = File):
    files_list = [file_1,file_2]

    temp_doc_pdf = None  
    temp_pdf = None  

    archivos_temp = []

    for file in files_list:
        if file.filename.lower().endswith(".docx"):
            temp_doc_pdf = await TempleFileHannder.create_temple_doc_to_pdf(file=file.file)
            archivos_temp.append(temp_doc_pdf)
        elif file.filename.lower().endswith(".pdf"):
            temp_pdf = await TempleFileHannder.create_templife_pdf(file=file.file)
            archivos_temp.append(temp_pdf)

    temp_md_1 = await TempleFileHannder.make_md_file(temp_path=temp_pdf)

    archivos_temp.append(temp_md_1)

    temp_md_2 = await TempleFileHannder.make_md_file(temp_path=temp_doc_pdf)

    archivos_temp.append(temp_md_2)

    prompt_template = PromptTemplate(
        input_variables=["md_isp","md_adium"],
        template= await read_md("promptSystem_5.md")
    )

    with open(file=temp_md_1, mode="r", encoding="utf-8") as file:
        md_data_1 = file.read()

    llm_clean_md = await chat_completation_doc(md_text=md_data_1)

    await create_docx(parragragh=llm_clean_md)

    md_colors = await procesar_prospecto(md_text=md_data_1)


    with open(file=temp_md_2, mode="r", encoding="utf-8") as file:
        md_data_2 = file.read()


    final_prompt = prompt_template.format(
        md_isp=md_colors,
        md_adium=md_data_2,
    )

    await asyncio.sleep(0.2)  
    # Elimina todos los archivos temporales creados, si existen  
    # for temp_path in archivos_temp:  
    #     if temp_path and os.path.exists(temp_path):  
    #         try:  
    #             os.remove(temp_path)  
    #         except Exception as e:  
    #             print(f"Error al borrar {temp_path}: {str(e)}")  

    await TempleFileHannder.delete_templife(temp_doc_pdf)
    await TempleFileHannder.delete_templife(temp_pdf)
    await TempleFileHannder.delete_templife(temp_md_1)
    await TempleFileHannder.delete_templife(temp_md_2)

    response = await chat_comparation(prompt=final_prompt)
    
    #file_new = current_dir + f"/archivo_final_{id}.md"
    id = get_id()
    file_name = f"archivo_final_{id}.md"
    file_path = os.path.join('documents', file_name)
    try:
        with open(file_path, mode="w", encoding="utf-8") as f:
            f.write(response)
            return {"response_comparation":"ARCHIVO MARKDOWN LISTO!"}
    except Exception as e:
        raise Exception(e)


    

   